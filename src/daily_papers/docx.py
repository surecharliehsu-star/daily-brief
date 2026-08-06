import json
from datetime import datetime, date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent.parent
OUTPUT_DIR = ROOT / "outputs"


def _set_run_font(run, name_cn="仿宋", name_en="Times New Roman", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.name = name_en
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), name_cn)
    if color:
        run.font.color.rgb = color


def _add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_run_font(run, name_cn="黑体", size=16, bold=True)
    elif level == 1:
        run = p.add_run(text)
        _set_run_font(run, name_cn="黑体", size=14, bold=True)
    elif level == 2:
        run = p.add_run(text)
        _set_run_font(run, name_cn="楷体", size=12, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_body(doc, text, size=12, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def _add_hyperlink(paragraph, text, url):
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def _to_serial(num):
    cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if num <= 10:
        return cn_nums[num - 1]
    return str(num)


def _load_papers(day_dir: Path) -> list[dict]:
    path = day_dir / "papers.json"
    if not path.exists():
        return []
    return json.loads(path.read_text(encoding="utf-8"))


def generate_docx(day_dir: Path) -> str | None:
    papers = _load_papers(day_dir)
    if not papers:
        return None

    papers.sort(key=lambda p: (p.get("source", ""), not p.get("is_monetary", False), p.get("published", "") or ""))
    from collections import OrderedDict
    grouped = OrderedDict()
    for p in papers:
        src = p.get("source", "其他")
        grouped.setdefault(src, []).append(p)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    today = date.today()
    date_str = today.strftime("%Y年%m月%d日")

    _add_heading(doc, f"国际货币政策文献每日简报", level=0)
    _add_body(doc, f"编制日期：{date_str}    覆盖机构：{len(grouped)}家    文献总数：{len(papers)}篇", size=11, indent=False)
    _add_body(doc, "", size=6, indent=False)

    seq = 1
    for src, src_papers in grouped.items():
        _add_heading(doc, f"{src}（{len(src_papers)}篇）", level=1)
        for p in src_papers:
            title_zh = (p.get("title_zh", "") or "").strip()
            title_en = (p.get("title", "") or "").strip()
            authors = p.get("authors", [])
            author_str = ", ".join(authors[:3]) if authors else ""
            if len(authors) > 3:
                author_str += " et al."
            pub = (p.get("published", "") or "")[:10]
            url = p.get("url", "") or ""
            is_mon = p.get("is_monetary", False)

            line = f"{seq}. "
            if title_zh:
                line += title_zh
            else:
                line += title_en

            meta_parts = [f"    {title_en}"]
            if author_str:
                meta_parts.append(author_str)
            if pub:
                meta_parts.append(pub)

            p_para = doc.add_paragraph()
            p_para.paragraph_format.line_spacing = Pt(22)
            p_para.paragraph_format.space_before = Pt(3)
            p_para.paragraph_format.space_after = Pt(1)

            run = p_para.add_run(line)
            _set_run_font(run, size=12, bold=True)

            if url:
                _add_hyperlink(p_para, " [原文]", url)

            meta_text = " · ".join(meta_parts)
            _add_body(doc, meta_text, size=11, indent=False)

            if is_mon:
                abstract_zh = p.get("abstract_zh", "") or ""
                if abstract_zh:
                    abs_para = doc.add_paragraph()
                    abs_para.paragraph_format.line_spacing = Pt(22)
                    abs_para.paragraph_format.space_before = Pt(1)
                    abs_para.paragraph_format.space_after = Pt(4)
                    abs_para.paragraph_format.first_line_indent = Cm(0.74)
                    run = abs_para.add_run(abstract_zh)
                    _set_run_font(run, size=11)
                    color_run = abs_para.add_run(" [货币政策]")
                    _set_run_font(color_run, size=10, bold=True, color=RGBColor(0xCF, 0x13, 0x22))
                else:
                    abs_para = doc.add_paragraph()
                    abs_para.paragraph_format.line_spacing = Pt(22)
                    abs_para.paragraph_format.space_before = Pt(1)
                    abs_para.paragraph_format.space_after = Pt(4)
                    run = abs_para.add_run("（摘要暂缺，次日自动补齐）")
                    _set_run_font(run, size=11, color=RGBColor(0x99, 0x99, 0x99))

            seq += 1

    _add_body(doc, "", size=6, indent=False)
    _add_body(doc, f"自动生成 · {today.strftime('%Y-%m-%d %H:%M')}", size=10, indent=False, bold=False)

    brief_path = day_dir / "brief.docx"
    doc.save(str(brief_path))
    return str(brief_path)
