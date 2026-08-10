"""全文翻译的 Word/PDF 文档生成（含中文排版）。"""
from __future__ import annotations

from datetime import date
from pathlib import Path


def _gen_docx(path: Path, title: str, paper: dict, translated: str, original: str) -> None:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    def font(run, size=12, bold=False, cn="仿宋", en="Times New Roman"):
        run.font.size = Pt(size)
        run.font.name = en
        run.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(title), size=16, bold=True, cn="黑体")

    meta = f"{date.today().strftime('%Y-%m-%d')} · {paper.get('source', '')}"
    if paper.get("published"):
        meta += f" · 发布于 {paper.get('published', '')[:10]}"
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(m.add_run(meta), size=11)

    for para in translated.split("\n\n"):
        if not para.strip():
            continue
        body = doc.add_paragraph()
        body.paragraph_format.line_spacing = Pt(24)
        body.paragraph_format.first_line_indent = Cm(0.74)
        font(body.add_run(para.strip()), size=12)

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(12)
    font(foot.add_run(f"技术说明：本译文由 AI 自动翻译，仅供参考；原文见 {paper.get('url') or paper.get('pdf_url', '')}"), size=9)

    doc.save(str(path))


def _gen_pdf(path: Path, title: str, paper: dict, translated: str) -> bool:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

        doc = BaseDocTemplate(str(path), pagesize=A4,
                              leftMargin=28 * mm, rightMargin=28 * mm,
                              topMargin=25 * mm, bottomMargin=25 * mm)
        frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
        doc.addPageTemplates([PageTemplate(id="all", frames=[frame])])

        title_style = ParagraphStyle("t", fontName="STSong-Light", fontSize=16, leading=22,
                                     alignment=1, spaceAfter=6)
        meta_style = ParagraphStyle("m", fontName="STSong-Light", fontSize=10, leading=14,
                                    alignment=1, spaceAfter=14, textColor="#666666")
        body_style = ParagraphStyle("b", fontName="STSong-Light", fontSize=12, leading=20,
                                    firstLineIndent=24, spaceAfter=6, wordWrap="CJK")
        foot_style = ParagraphStyle("f", fontName="STSong-Light", fontSize=9, leading=12,
                                    textColor="#999999", spaceBefore=12)

        from html import escape
        story = []
        story.append(Paragraph(escape(title), title_style))
        meta = f"{date.today().strftime('%Y-%m-%d')} · {escape(str(paper.get('source', '')))}"
        if paper.get("published"):
            meta += f" · 发布于 {escape(str(paper.get('published', ''))[:10])}"
        story.append(Paragraph(meta, meta_style))
        for para in translated.split("\n\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(escape(para).replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 10))
        note = f"技术说明：本译文由 AI 自动翻译，仅供参考；原文见 {escape(str(paper.get('url') or paper.get('pdf_url', '')))}"
        story.append(Paragraph(note, foot_style))

        doc.build(story)
        return True
    except Exception as e:
        print(f"[PDF] reportlab failed: {e}", flush=True)
        return False


def generate_fulltext_docs(paper: dict, translated: str, output_dir: Path) -> list[Path]:
    """用已翻译好的中文文本生成 .docx 与 .pdf，返回文件路径列表。"""
    title = paper.get("title_zh") or paper.get("title")

    safe_title = "".join(c for c in title if c not in '/\\:*?"<>|').strip()[:60] or "translated"
    day = output_dir / "".join(str(date.today()).split("-"))
    out_dir = day / "fulltext_zh"
    out_dir.mkdir(parents=True, exist_ok=True)

    files = []
    try:
        docx_path = out_dir / f"{safe_title}.docx"
        _gen_docx(docx_path, title, paper, translated, translated)
        files.append(docx_path)
    except Exception as e:
        print(f"[DOCX] {safe_title}.docx failed: {e}", flush=True)

    try:
        pdf_path = out_dir / f"{safe_title}.pdf"
        if _gen_pdf(pdf_path, title, paper, translated):
            files.append(pdf_path)
    except Exception as e:
        print(f"[PDF] failed: {e}", flush=True)

    return files